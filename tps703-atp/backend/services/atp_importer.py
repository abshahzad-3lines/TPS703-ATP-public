"""Phase 10 — DOCX / PDF importer for customer-uploaded ATP documents.

Two layers:
1. ``extract_text_from_docx`` / ``extract_text_from_pdf`` — deterministic
   text extraction using python-docx and pdfplumber. Returns the full
   plain-text contents (paragraphs + table cells, joined with newlines).
2. ``heuristic_extract_steps`` — best-effort regex-based step extractor
   that splits on "Step N", "Step N.", "N.", etc. headings and packages
   the surrounding text as ``instructions``. Engineers refine the result
   in the authoring UI (or hand it to the Wave-5 AI enricher).

The heuristic is intentionally conservative: it never guesses step_type,
instrument, frequency or limits — those carry safety implications and
must be confirmed by a human or by the AI-with-review pass.
"""

from __future__ import annotations

import io
import re
from typing import Iterable


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text_from_docx(data: bytes) -> str:
    """Return paragraph + table-cell text from a .docx blob."""
    try:
        from docx import Document
    except ImportError as e:  # noqa: F401
        raise RuntimeError(
            "python-docx is not installed. Run `pip install python-docx`."
        ) from e

    doc = Document(io.BytesIO(data))
    chunks: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells)
            if row_text.replace("|", "").strip():
                chunks.append(row_text)
    return "\n".join(chunks)


def extract_text_from_pdf(data: bytes) -> str:
    """Return text from every page of a PDF, joined with form-feeds."""
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError(
            "pdfplumber is not installed. Run `pip install pdfplumber`."
        ) from e

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    # \f delimiter so downstream consumers can split per page if they want.
    return "\f".join(pages)


# ---------------------------------------------------------------------------
# Heuristic step splitter
# ---------------------------------------------------------------------------


_STEP_HEAD_RE = re.compile(
    r"""(?ix)
    ^\s*
    (?:step\s+)?          # optional 'step '
    (\d+)                 # capture step number
    [\.\):]\s+            # one of . ) :
    (.{4,120}?)           # short heading (4-120 chars)
    \s*$
    """,
    re.MULTILINE,
)


def heuristic_extract_steps(text: str) -> list[dict]:
    """Return a list of ``{step_number, name, step_type, instructions}`` dicts.

    ``step_type`` is always ``'manual_record'`` so the validator does not
    error before the engineer or AI enriches the rows. ``instructions``
    contains the captured body between this heading and the next.
    """
    matches = list(_STEP_HEAD_RE.finditer(text))
    if not matches:
        return []

    out: list[dict] = []
    for i, m in enumerate(matches):
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[body_start:body_end].strip()
        # First non-empty line of the body becomes the additional title
        out.append({
            "step_number": int(m.group(1)),
            "name": _truncate(m.group(2).strip(), 120),
            "step_type": "manual_record",
            "instructions": _truncate(body, 4000) or None,
        })

    # Re-number 1..N to honour the schema constraint, preserving order.
    out.sort(key=lambda d: d["step_number"])
    for idx, row in enumerate(out, start=1):
        row["step_number"] = idx
    return out


def _truncate(s: str, n: int) -> str:
    if len(s) <= n:
        return s
    return s[: n - 1].rstrip() + "…"


# ---------------------------------------------------------------------------
# Metadata heuristics — pull a code + name out of the document header
# ---------------------------------------------------------------------------


_CODE_RE = re.compile(
    r"(?im)^\s*(?:document\s+(?:number|no\.?|#)|drawing\s+no\.?|atp\s+code|procedure\s+code)\s*[:\-]\s*([A-Z0-9_\-]+)"
)
_NAME_RE = re.compile(
    r"(?im)^\s*(?:title|procedure\s+name|test\s+name|atp\s+name)\s*[:\-]\s*(.{4,120})"
)


def guess_metadata(text: str) -> dict:
    """Return ``{code, name}`` best-guesses; either may be missing."""
    out: dict = {}
    m = _CODE_RE.search(text)
    if m:
        out["code"] = m.group(1).strip().upper()
    m = _NAME_RE.search(text)
    if m:
        out["name"] = _truncate(m.group(1).strip(), 120)
    return out
